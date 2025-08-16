import streamlit as st
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from pymongo import MongoClient
import datetime
import base64
from io import BytesIO
import pymongo
import fitz
from docx import Document
from bson import ObjectId
import re
import spacy
from pyresparser import ResumeParser
import tempfile
from rapidfuzz import fuzz
import smtplib
from email.message import EmailMessage

def send_application_email(to_email, role, company):
    sender_email = "sanjaychandiran@gmail.com"  # Replace with your sender email
    app_password = "petd tstk eagj ijqv"     # Replace with your Gmail App Password

    msg = EmailMessage()
    msg['Subject'] = f"Application Submitted: {role} at {company}"
    msg['From'] = sender_email
    msg['To'] = to_email
    msg.set_content(f"""\
Dear Candidate,

You have successfully applied for the role of {role} at {company}.

Thank you for using our platform!

Best regards,
Internship Portal Team
""")

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(sender_email, app_password)
            smtp.send_message(msg)
        return True
    except Exception as e:
        print("Error sending email:", e)
        return False


# MongoDB Atlas connection
client = MongoClient("mongodb+srv://Sanjay:Sanjay@cluster0.ft1o2.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
custom_nlp = spacy.load("en_core_web_sm")
# Database and collections
db = client["Sanjay"]  # Use your actual DB name
users_collection = db["Kavin"]  # This collection will store user profiles
internships_collection = db["ragul"]  # For storing internship data
profiles_collection = db["nithish"]  # For storing detailed user profiles
# Add this line with your existing collections
apprenticeships_collection = db["ApprentNtern"]  # New collection for apprenticeship listings



# Session state init
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "show_profile" not in st.session_state:
    st.session_state.show_profile = False
if "show_signup" not in st.session_state:
    st.session_state.show_signup = False
if "page" not in st.session_state:
    st.session_state.page = "login"  # can be login, internexus, main

# States of India
indian_states = [
    "Andhra Pradesh", "Arunachal Pradesh", "Assam", "Bihar", "Chhattisgarh", "Goa", "Gujarat",
    "Haryana", "Himachal Pradesh", "Jharkhand", "Karnataka", "Kerala", "Madhya Pradesh",
    "Maharashtra", "Manipur", "Meghalaya", "Mizoram", "Nagaland", "Odisha", "Punjab", "Rajasthan",
    "Sikkim", "Tamil Nadu", "Telangana", "Tripura", "Uttar Pradesh", "Uttarakhand", "West Bengal"
]
indian_cities = [
    # Andhra Pradesh
    "Visakhapatnam", "Vijayawada", "Guntur", "Nellore", "Tirupati", "Rajahmundry", "Kakinada",
    # Arunachal Pradesh
    "Itanagar", "Tawang", "Ziro", "Bomdila", "Pasighat",
    # Assam
    "Guwahati", "Dibrugarh", "Jorhat", "Silchar", "Nagaon", "Tezpur",
    # Bihar
    "Patna", "Gaya", "Bhagalpur", "Muzaffarpur", "Munger", "Darbhanga",
    # Chhattisgarh
    "Raipur", "Bilaspur", "Korba", "Durg", "Raigarh",
    # Goa
    "Panaji", "Vasco da Gama", "Margao", "Mapusa",
    # Gujarat
    "Ahmedabad", "Surat", "Vadodara", "Rajkot", "Bhavnagar", "Junagadh", "Anand",
    # Haryana
    "Chandigarh", "Faridabad", "Gurugram", "Ambala", "Hisar", "Panipat",
    # Himachal Pradesh
    "Shimla", "Manali", "Dharamshala", "Kullu", "Solan",
    # Jharkhand
    "Ranchi", "Jamshedpur", "Dhanbad", "Deoghar", "Hazaribagh",
    # Karnataka
    "Bengaluru", "Mysuru", "Hubballi", "Mangalore", "Belagavi", "Tumakuru",
    # Kerala
    "Thiruvananthapuram", "Kochi", "Kozhikode", "Kottayam", "Thrissur", "Malappuram",
    # Madhya Pradesh
    "Bhopal", "Indore", "Gwalior", "Jabalpur", "Ujjain", "Sagar",
    # Maharashtra
    "Mumbai", "Pune", "Nagpur", "Nashik", "Aurangabad", "Thane", "Solapur",
    # Manipur
    "Imphal", "Thoubal", "Churachandpur", "Bishnupur",
    # Meghalaya
    "Shillong", "Tura", "Jowai", "Nongstoin",
    # Mizoram
    "Aizawl", "Lunglei", "Champhai",
    # Nagaland
    "Kohima", "Dimapur", "Mokokchung",
    # Odisha
    "Bhubaneswar", "Cuttack", "Rourkela", "Berhampur", "Sambalpur",
    # Punjab
    "Chandigarh", "Amritsar", "Ludhiana", "Jalandhar", "Patiala",
    # Rajasthan
    "Jaipur", "Udaipur", "Jodhpur", "Kota", "Ajmer", "Bikaner",
    # Sikkim
    "Gangtok", "Mangan", "Namchi",
    # Tamil Nadu
    "Chennai", "Coimbatore", "Madurai", "Salem", "Trichy", "Tirunelveli",
    # Telangana
    "Hyderabad", "Warangal", "Khammam", "Nizamabad",
    # Tripura
    "Agartala", "Udaipur", "Dharmanagar",
    # Uttar Pradesh
    "Lucknow", "Kanpur", "Varanasi", "Agra", "Allahabad", "Ghaziabad",
    # Uttarakhand
    "Dehradun", "Haridwar", "Nainital", "Rishikesh", "Mussoorie",
    # West Bengal
    "Kolkata", "Darjeeling", "Howrah", "Siliguri", "Asansol", "Durgapur",
    "New Delhi", "Delhi", "Rohini", "Dwarka", "Connaught Place", "Chandigarh",
    "Puducherry", "Karaikal", "Mahe", "Yanam",
    "Leh", "Kargil", "Diskit", "Nubra", "Srinagar", "Jammu", "Anantnag", 
    "Baramulla", "Udhampur", "Port Blair", "Diglipur", "Mayabunder", "Rangat", "Kavaratti",
    "Agatti", "Minicoy", "Kalpeni", "Silvassa", "Daman", "Diu"

]


if "users" not in st.session_state:
    st.session_state.users = []
if "user_email" not in st.session_state:
    st.session_state.user_email = ""
  

# --- UI FUNCTIONS ---
def show_login():
    st.markdown(
        """
        <style>
        .stApp {
            background-image: url("https://images.unsplash.com/photo-1517511620798-cec17d428bc0");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
            color:black
        }
        label {
            color:black !important;
            font-weight:bold;
        }
        input, textarea {
            background-color: white !important;
            color: black !important;
            border: 1px solid #ccc !important;
            border-radius: 5px !important;
        }
        .stButton > button {
            background-color: #007BFF !important; /* Bootstrap blue */
            color: white !important;
            font-weight: bold;
            border: none;
            padding: 10px 20px;
            border-radius: 8px;
            transition: background-color 0.3s ease;
        }
        .stButton > button:hover {
            background-color: #0056b3 !important; /* Darker blue on hover */
            color: white !important;
        }
        
        </style>
        """,
        unsafe_allow_html=True
    )
    
    st.title("Internship and Apprenticeships Portal")
    login_input = st.text_input("Email or Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if "@" in login_input and not login_input.endswith("@gmail.com"):
            st.warning("Please enter a valid Gmail address ending with '@gmail.com'.")
        else:   
        # Check MongoDB for matching user credentials
            matched_user = users_collection.find_one({
                "$or": [{"email": login_input}, {"username": login_input}],
                "password": password
            })

            if matched_user:
                st.session_state.logged_in = True
                st.session_state.page = "internexus"
                st.session_state.show_signup = False
                st.session_state.user_email = matched_user.get("email","")
                st.success(f"Welcome, {matched_user['first_name']}!")  # Store user's email for later
                st.rerun()
            else:
                st.warning("Invalid email/username or password.")


    if st.button("Sign Up"):
        st.session_state.show_signup = True
        st.rerun()
def show_signup():
    st.markdown(
        """
        <style>
        .stApp {
            background-image: url("https://static.vecteezy.com/system/resources/thumbnails/002/375/040/small_2x/modern-white-background-free-vector.jpg");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
            color: black;
        }

        label {
            color: black !important;
            font-weight:bold;
        }

        input, textarea {
            background-color: white !important;
            color: black !important;
            
        }
        div.stCheckbox > label > div[data-baseweb="checkbox"] > div > span {
            color: black !important;
            font-weight: bold;
        }

        .stButton > button {
            background-color: #007BFF !important;
            color: white !important;
            font-weight: bold;
            border-radius: 8px;
            padding: 8px 16px;
        }

        .stButton > button:hover {
            background-color: #0056b3 !important;
            color: white !important;
        }
         div.stAlert > div > div {
            color: black !important;
            font-weight: bold;
        }
        div.stAlert-warning {
            background-color: #f0f0f0 !important;
        }
        div.stAlert-error {
            background-color: #f0f0f0 !important;
        }
        div.stAlert-info {
            background-color: #f0f0f0 !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.title("Sign Up into Portal")
    
    first_name = st.text_input("First Name")
    last_name = st.text_input("Last Name")
    username = st.text_input("Username")
    email = st.text_input("Email ID")
    password = st.text_input("Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")
    captcha = st.checkbox("")
    st.markdown("<span style='color: black; font-weight: bold;'>I am not a robot</span>", unsafe_allow_html=True)

    if st.button("Register"):
        if not (first_name and last_name and username and email and password and confirm_password):
            st.warning("Please fill all fields.")
        elif password != confirm_password:
            st.warning("Passwords do not match.")
        elif not captcha:
            st.warning("Please verify that you're not a robot.")
        elif users_collection.find_one({"$or": [{"email": email}, {"username": username}]}):
            st.warning("User with this email or username already exists.")
        else:
            # Save new user to MongoDB
            user_data = {
                "first_name": first_name,
                "last_name": last_name,
                "username": username,
                "email": email,
                "password": password
            }
            users_collection.insert_one(user_data)
            st.success("Registration Successful! Please login.")
            st.session_state.show_signup = False
            st.rerun()

def show_internexus():
    st.markdown(
        """
        <style>
        .stApp {
            background-image: url("https://images.unsplash.com/photo-1503676260728-1c00da094a0b?auto=format&fit=crop&w=1470&q=80");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }

        .stButton > button {
            background-color: #007BFF !important;
            color: white !important;
            font-weight: bold;
            padding: 8px 20px;
            border-radius: 8px;
            border: none;
            transition: background-color 0.3s ease;
        }

        .stButton > button:hover {
            background-color: #0056b3 !important;
            color: white !important;
        }

        .internexus-box {
            background: rgba(255, 255, 255, 0.25);
            backdrop-filter: blur(12px);
            -webkit-backdrop-filter: blur(12px);
            border-radius: 15px;
            padding: 30px;
            box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.2);
            border: 1px solid rgba(255, 255, 255, 0.3);
        }

        h1.internexus-title {
            text-align: center;
            font-size: 3rem;
            font-weight: 900;
            margin-top: 20px;
            margin-bottom: 40px;
            color: black;
        }

        .internexus-text {
            color: black;
            font-size: 1rem;
            line-height: 1.6;
            text-align: center;
        }

        .internexus-heading {
            font-size: 1.6rem;
            font-weight: bold;
            text-align: center;
            color: black;
            margin-bottom: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Title
    st.markdown('<h1 class="internexus-title">ApprentNtern</h1>', unsafe_allow_html=True)

    # Two side-by-side columns
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<div class="internexus-box">', unsafe_allow_html=True)
        st.markdown('<div class="internexus-heading">Internship</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="internexus-text">Explore a wide range of internship opportunities tailored for students and recent graduates. Get access to government and private internships in software, hardware, AI, ML, cloud, and more.</div>',
            unsafe_allow_html=True
        )
        st.write("")
        if st.button("Go to Internship Platform", key="go_internship"):
            st.session_state.page = "main"
            st.session_state.internship_mode = "Internship"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="internexus-box">', unsafe_allow_html=True)
        st.markdown('<div class="internexus-heading">Apprenticeship</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="internexus-text">Join hands-on apprenticeship programs to gain real-world experience while learning. Ideal for those seeking to combine education with industry work in software, hardware, or research fields.</div>',
            unsafe_allow_html=True
        )
        st.write("")
        if st.button("Go to Apprenticeship Platform", key="go_apprenticeship"):
            st.session_state.page = "main"
            st.session_state.internship_mode = "Apprenticeship"
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

import re
import pdfplumber

def parse_resume(uploaded_file):
    extracted_info = {
        "Name": "Not detected",
        "Email": "Not detected",
        "Phone": "Not detected",
        "Skills": []
    }

    # Extract text using pdfplumber
    with pdfplumber.open(uploaded_file) as pdf:
        text = "\n".join([page.extract_text() or "" for page in pdf.pages])

    # Extract email
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    if email_match:
        extracted_info["Email"] = email_match.group()

    # Extract phone number
    phone_match = re.search(r"(\+?\d{1,3})?[\s-]?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}", text)
    if phone_match:
        extracted_info["Phone"] = phone_match.group()

    # Try to extract name (naive approach: first line or before email)
    lines = text.strip().split("\n")
    if lines:
        first_line = lines[0].strip()
        if len(first_line.split()) <= 4 and not any(x in first_line.lower() for x in ["email", "phone", "@"]):
            extracted_info["Name"] = first_line

    # Extract skills using simple matching
    common_skills = [
        "Python", "Java", "C++", "SQL", "Machine Learning", "Deep Learning", "NLP", "Streamlit",
        "HTML", "CSS", "JavaScript", "Django", "Flask", "React", "Node.js", "AWS", "Docker", "Git"
    ]
    skills_found = [skill for skill in common_skills if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
    extracted_info["Skills"] = list(set(skills_found))

    return text, extracted_info


def show_profile():
    st.markdown("""
    <style>
    /* Expander header title styling */
    .st-expanderHeader p {
        color: #000000 !important;
        font-weight: 900 !important;
        font-size: 18px !important;
    }

    /* Input field labels inside expander */
    .st-expander .stTextInput label,
    .st-expander .stTextArea label,
    .st-expander .stSelectbox label,
    .st-expander label {
        color: #000000 !important;
        font-weight: 800 !important;
        font-size: 16px !important;
    }

    /* Optional: make placeholder text darker for better visibility */
    .st-expander input::placeholder {
        color: #333333 !important;
    }

    /* Optional: make expander content text black */
    .st-expander p, .st-expander div {
        color: #000000 !important;
    }

    /* Selection color */
    ::selection {
        background-color: #333333 !important;
        color: #ffffff !important;
    }
    ::-moz-selection {
        background-color: #333333 !important;
        color: #ffffff !important;
    }
    </style>
    """, unsafe_allow_html=True)

   


    st.title("User Profile")
    st.subheader("Personal Details")
    name = st.text_input("Full Name")
    email = st.text_input("Email ID")
    phone = st.text_input("Phone Number")  # ✅ Phone field
    dob = st.date_input(
    "Date of Birth",
    min_value=datetime.date(1990, 1, 1),
    max_value=datetime.date(2050, 12, 31)
)
    gender = st.radio("Gender", ["Male", "Female", "Other"])
    location = st.selectbox("Select Location", indian_states)
    github = st.text_input("GitHub Profile URL")
    linkedin = st.text_input("LinkedIn Profile URL")

    st.subheader("Education Details")
    university = st.text_input("University Name")
    degree = st.text_input("Degree")
    branch = st.text_input("Branch")
    current_year = st.selectbox("Current Year", ["1st Year", "2nd Year", "3rd Year", "4th Year"])
    cgpa = st.text_input("Current CGPA")
    expected_graduation_month = st.selectbox("Expected Graduation Month", [
        "January", "February", "March", "April", "May", "June", 
        "July", "August", "September", "October", "November", "December"])
    expected_graduation_year = st.selectbox("Expected Graduation Year", [str(i) for i in range(2024, 2035)])

    st.subheader("Certifications")
    if "certifications" not in st.session_state:
        st.session_state.certifications = []

    for i, cert in enumerate(st.session_state.certifications):
        with st.expander("Certification Details"):
            cert["name"] = st.text_input("Certificate Name", cert["name"], key=f"cert_name_{i}")
            cert["id"] = st.text_input("Certificate ID", cert["id"], key=f"cert_id_{i}")
            cert["url"] = st.text_input("Certificate URL", cert["url"], key=f"cert_url_{i}")
            cert["start_date"] = st.date_input("Start Date", cert["start_date"], key=f"cert_start_{i}")
            cert["skills"] = st.text_input("Skills from Certification", cert["skills"], key=f"cert_skills_{i}")
            if st.button("Remove Certification", key=f"remove_cert_{i}"):
                del st.session_state.certifications[i]
                st.rerun()

    if st.button("+ Add Certification",key="add_cert_btn"):
        st.session_state.certifications.append({
            "name": "", "id": "", "url": "", "start_date": None,
            "end_date": None, "skills": ""
        })
        st.rerun()

    st.subheader("Skills")

    hard_skill_options = [
    "Python", "Java", "C++", "C", "JavaScript", "React", "Node.js", "SQL", "Django", "Flask",
    "TensorFlow", "PyTorch", "Keras", "Pandas", "NumPy", "Matplotlib", "Git", "Docker", "AWS",
    "Azure", "Linux", "MongoDB", "Firebase", "Power BI", "Tableau", "HTML", "CSS", "C#", "Go", "Rust"
    ]

    soft_skill_options = [
    "Communication", "Teamwork", "Leadership", "Time Management", "Adaptability",
    "Problem Solving", "Creativity", "Critical Thinking", "Work Ethic", "Emotional Intelligence",
    "Decision Making", "Conflict Resolution", "Public Speaking", "Active Listening"
    ]

    hard_skills = st.multiselect("Select Hard Skills:", options=hard_skill_options)
    soft_skills = st.multiselect("Select Soft Skills:", options=soft_skill_options)

    st.subheader("Projects")
    if "projects" not in st.session_state:
        st.session_state.projects = []

    for i, project in enumerate(st.session_state.projects):
        with st.expander("Project Details"):
            project["title"] = st.text_input("Project Title", project["title"], key=f"proj_title_{i}")
            project["description"] = st.text_area("Project Description", project["description"], key=f"proj_desc_{i}")
            project["technologies"] = st.text_input("Technologies Used", project["technologies"], key=f"proj_tech_{i}")
            project["github"] = st.text_input("Project GitHub URL", project["github"], key=f"proj_github_{i}")
            if st.button("Remove Project", key=f"remove_proj_{i}"):
                del st.session_state.projects[i]
                st.rerun()

    if st.button("+ Add Project",key="add_proj_btn"):
        st.session_state.projects.append({
            "title": "", "description": "", "technologies": "", "github": ""
        })
        st.rerun()
    st.subheader("Academic Achievements")

    if "academic_achievements" not in st.session_state:
        st.session_state.academic_achievements = []

    for i, achievement in enumerate(st.session_state.academic_achievements):
        with st.expander(f"Achievement {i+1}"):
            achievement["title"] = st.text_input("Title", achievement.get("title", ""), key=f"ach_title_{i}")
            achievement["description"] = st.text_area("Description", achievement.get("description", ""), key=f"ach_desc_{i}")
            achievement["year"] = st.text_input("Year/Awarded On", achievement.get("year", ""), key=f"ach_year_{i}")

            if st.button("Remove Achievement", key=f"remove_ach_{i}"):
                del st.session_state.academic_achievements[i]
                st.rerun()

    if st.button("+ Add Academic Achievement",key="add_ach_btn"):
        st.session_state.academic_achievements.append({
            "title": "", "description": "", "year": ""
        })
        st.rerun()

         # Show Work Experience only in Apprenticeship mode
        if st.button("+ Add Work Experience",key="add_work_btn"):
            st.session_state.work_experience.append({
                "company": "", "role": "", "duration": "", "description": ""
            })
            st.rerun()

    st.subheader("Languages Known")
    languages_known = st.multiselect("Select Languages You Know:", [
        "English", "Hindi", "Tamil", "Telugu", "Kannada", "Malayalam", "Marathi", "Gujarati", 
        "Bengali", "Punjabi", "Odia"])

    st.subheader("Profile Summary")
    profile_summary = st.text_area("Write a short summary about yourself")
   
    required_skills = set([
    "Python", "SQL", "Machine Learning", "Data Analysis", "Communication",
    "Pandas", "scikit-learn", "Teamwork", "Git", "Problem Solving"
    ])

    def extract_text_from_pdf(file_path):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text

    def extract_text_from_docx(file_path):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def extract_email(text):
        match = re.search(r'[\w\.-]+@[\w\.-]+', text)
        return match.group(0) if match else "N/A"

    def extract_phone(text):
        match = re.search(r'(\+91[\-\s]?)?[0]?[789]\d{9}', text)
        return match.group(0) if match else "N/A"

    def extract_name(text):
        lines = text.split('\n')
        for line in lines:
            if line.strip() and len(line.strip().split()) <= 4:
                return line.strip()
        return "N/A"

    def extract_skills(text):
        found = []
        for skill in required_skills:
            if skill.lower() in text.lower():
                found.append(skill)
        return found
    def parse_resume(file_bytes, file_type):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf" if file_type == "application/pdf" else ".docx") as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        if file_type == "application/pdf":
            text = extract_text_from_pdf(tmp_path)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(tmp_path)
        else:
            text = file_bytes.decode("utf-8")

    # Basic cleaning and formatting
        extracted_info = {
            "Name": extract_name(text),
            "Email": extract_email(text),
            "Phone": extract_phone(text),
            "Skills": extract_skills(text),
            "Experience": "Yes" if "experience" in text.lower() else "No",
            "Education": "Yes" if "education" in text.lower() else "No",
            "Certifications": "Yes" if "certification" in text.lower() else "No",
            "Projects": "Yes" if "project" in text.lower() else "No"
        }
        matched = set(skill.lower() for skill in extracted_info["Skills"])
        required = set(skill.lower() for skill in required_skills)
        score = int((len(matched & required) / len(required)) * 100)

        return text, extracted_info, score
    resume_file = st.file_uploader("Upload Your Resume",type=["pdf","docx","txt"])    
    if resume_file is not None:
        if resume_file.size > 300 * 1024:
            st.error("❌ File size exceeds 300KB.")
        else:
            resume_bytes = resume_file.read()
            resume_base64 = base64.b64encode(resume_bytes).decode("utf-8")

        # 📄 Resume preview
            if resume_file.type == "application/pdf":
                pdf_display = f'<iframe src="data:application/pdf;base64,{resume_base64}" width="700" height="1000" type="application/pdf"></iframe>'
                st.markdown(pdf_display, unsafe_allow_html=True)

            elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                def read_docx(file):
                    doc = Document(file)
                    return "\n".join([para.text for para in doc.paragraphs])
                resume_text = read_docx(resume_file)
                st.text_area("📄 DOCX Resume Preview", resume_text, height=400)

            elif resume_file.type == "text/plain":
                resume_text = resume_bytes.decode("utf-8")
                st.text_area("📄 TXT Resume Preview", resume_text, height=400)

        # ⛏️ Resume parsing
        # Call parser
            resume_text, extracted_info, overall_score = parse_resume(resume_bytes, resume_file.type)

        # Match analysis
            found_skills = set(skill.lower() for skill in extracted_info.get("Skills", []))
            matched_skills = found_skills.intersection(skill.lower() for skill in required_skills)
            missing_skills = set(skill.lower() for skill in required_skills) - found_skills
            match_percentage = int((len(matched_skills) / len(required_skills)) * 100)

        # Display results
            with st.expander("📊 ATS Resume Analysis"):
                st.markdown("### Extracted Information:")
                for key, value in extracted_info.items():
                    st.write(f"**{key}**: {', '.join(value) if isinstance(value,list)else value}")
                found_skills = set(skill.lower() for skill in extracted_info.get("Skills", []))
                matched_skills = found_skills.intersection(skill.lower() for skill in required_skills)
                missing_skills = set(skill.lower() for skill in required_skills) - found_skills

                st.markdown("### 🔍 Match Analysis:")
                st.markdown(f"**Matched Skills**: `{', '.join(matched_skills)}`")
                st.markdown(f"**Missing Skills**: `{', '.join(missing_skills)}`")
                st.progress(match_percentage)
                st.markdown(f"**Overall Match Score**: `{match_percentage}%`")

                if match_percentage < 50:
                    st.warning("⚠️ Resume match is low. Consider adding more relevant skills or experience.")
                elif match_percentage < 80:
                    st.info("ℹ️ Resume match is decent. Enhance with a few more skills.")
                else:
                    st.success("✅ Great match! Your resume aligns well with the role.")

            with st.expander("📄 Full Resume Text (Optional)"):
                st.text_area("Resume Content", resume_text, height=400)
        if st.button("Save Profile",key="save_profile_btn"): 
            if not email:
                st.error("❌ Email is required to save your profile.")
            else:        
                profile_data = {
                    "name": name,
                    "email": email,
                    "phone": phone,
                    "dob": dob.strftime("%Y-%m-%d"),
                    "gender": gender,
                    "location": location,
                    "github": github,
                    "linkedin": linkedin,
                    "education": {
                        "university": university,
                        "degree": degree,
                        "branch": branch,
                        "current_year": current_year,
                        "cgpa": cgpa,
                            "expected_graduation": {
                                "month": expected_graduation_month,
                                "year": expected_graduation_year,
                            },
                        },
                        "certifications": st.session_state.certifications,
                        "skills": {
                            "hard_skills": hard_skills,
                            "soft_skills": soft_skills,
                        },
                        "projects": st.session_state.projects,
                        "academic_achievements": st.session_state.academic_achievements,
                        "languages_known": languages_known,
                        "profile_summary": profile_summary,
                        "resume":resume_base64 if resume_base64 else None,
                        "ats": {
                            "extracted_info": extracted_info,
                            "resume_text": resume_text,
                            "match_percentage": match_percentage,
                            "matched_skills": list(matched_skills),
                            "missing_skills": list(missing_skills),
                            "uploaded_filename": resume_file.name if resume_file else "No resume uploaded"
                        }
                    }
                result = profiles_collection.update_one(
                    {"email": email},           # Search by email
                    {"$set": profile_data},     # Set new or updated data
                    upsert=True                 # Insert if not found
                )
                if result.matched_count > 0:
                    st.success("✅ Profile updated successfully.")
                else:
                    st.success("✅ New profile created successfully.")
            # Assuming 'email' is a unique identifier
    if st.sidebar.button("Back to Dashboard"):
        st.session_state.show_profile = False
        st.rerun()
def show_main_page():
    st.markdown(
        """
        <style>
        section[data-testid="stSidebar"] .stButton > button {
            background-color: #007BFF !important;
            color: white !important;
            font-weight: bold !important;
            border-radius: 6px !important;
            padding: 0.5rem 1.2rem !important;
            margin-bottom: 10px;
        }
        input[type="text"],
        input[type="password"],
        input[type="email"],
        textarea {
            background-color: white !important;
            color: black !important;
            font-weight: 600;
            border-radius: 6px;
        }
        div[data-baseweb="select"] {
            background-color: white !important;
            color: black !important;
            font-weight: 600 !important;
        }


        section[data-testid="stSidebar"] .stButton > button:hover {
            background-color: #0056b3 !important;
            color: white !important;
        }
        section[data-testid="stSidebar"] label {
            color: white !important;
            font-weight: 600 !important;
        }
        /* Sidebar background and header */
        [data-testid="stSidebar"] {
            background-color: #1a1a1a;  /* Optional: dark background */
        }
        [data-testid="stSidebar"] h1,
        [data-testid="stSidebar"] h2,
        [data-testid="stSidebar"] h3,
        [data-testid="stSidebar"] h4,
        [data-testid="stSidebar"] h5,
        [data-testid="stSidebar"] h6 {
            color: white !important;
        }
        [data-testid="stSidebar"] label {
            color: white !important;
            font-weight: bold !important;
        }
        [data-testid="stSidebar"] .stSelectbox div,
        [data-testid="stSidebar"] .stMultiSelect div {
            color: white !important;
        }

        /* Sidebar header (e.g., st.sidebar.header("Dashboard")) */
        section[data-testid="stSidebar"] .css-1aumxhk {
            color: white !important;
            font-weight: 900 !important;
            font-size: 22px !important;
        }

        /* Sidebar label text (for selectboxes, etc.) */
        section[data-testid="stSidebar"] label {
            color: white !important;
            font-weight: 600 !important;
            font-size: 16px !important;
        }
        .stTextInput input,
        .stSelectbox div[data-baseweb="select"],
        .stMultiSelect div[data-baseweb="select"],
        .stDateInput input {
            background-color: white !important;
            color: black !important;
            border-radius: 6px !important;
        }
        .stSelectbox label, .stMultiSelect label, .stDateInput label {
            color: black !important;
            font-weight: 600 !important;
        }
        .stTextInput input {
            background-color: white !important;
            color: black !important;
            border: 1px solid #ccc !important;
            border-radius: 5px !important;
        }
        .stSelectbox div[data-baseweb="select"],
        .stMultiSelect div[data-baseweb="select"],
        .stDateInput input {
            background-color: white !important;
            color: black !important;
        }
        input::placeholder {
            color: #666 !important;
        }
        label {
            color: black !important;
            font-weight: 700 !important;
        }
       
        .stApp {
            background-image: url("https://www.google.com/url?sa=i&url=https%3A%2F%2Fwww.vecteezy.com%2Ffree-photos%2Fartificial-intelligence-background&psig=AOvVaw3qUdKf6fpVF2ZvenZanHQh&ust=1755367462875000&source=images&cd=vfe&opi=89978449&ved=0CBIQjRxqFwoTCMCdu8KzjY8DFQAAAAAdAAAAABAI");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #000000 !important;
            font-weight: 900 !important;
        }
        .stMarkdown p, .stMarkdown, .stExpander {
        color: #000000 !important;
        }
        .st-expanderContent, .st-expanderHeader {
        color: #000000 !important;
        font-weight: 600 !important;
        }
        .stExpander > summary {
        color: #000000 !important;
        font-weight: bold !important;
        font-size: 18px !important;
        }
        label, .css-1cpxqw2 {
            color: #000000 !important;
            font-weight: 800 !important;
            font-size: 16px !important;
        }
        .stSelectbox > div > div,
        .stMultiSelect > div > div {
            color: #000 !important;
            font-weight: 700 !important;
        }
        .st-expanderHeader p {
            color: #000000 !important;
            font-weight: 900 !important;
        }
        .stTextInput input::placeholder {
        color: #333333 !important;
        }
       
        
        </style>
        """,
        unsafe_allow_html=True
    )


    mode = st.session_state.get("internship_mode", "Internship")
    collection = internships_collection if mode == "Internship" else apprenticeships_collection
    selected_role = None
    stipend_filter = None
    st.title(f"AI-Powered {mode} Platform")
    st.sidebar.header("Dashboard")

    if st.sidebar.button("Profile"):
        st.session_state.show_profile = True
        st.rerun()

    # Internship type: Government or Private
    internship_type = st.sidebar.selectbox(f"Select {mode} Type:", ["Government", "Private"])

    # Internship category: Hardware or Software
    category = st.sidebar.selectbox("Select Category:", ["Hardware", "Software"])
    selected_role = None
    # Dynamic role options based on type and category
    if category:
        st.markdown(f"<h2 style='text-align: center;'> {internship_type} {category} Job Roles</h2>", unsafe_allow_html=True)
        roles_query = {
            "type": mode,
            "sector": internship_type,
            "category": category
        }
        roles_cursor = (internships_collection if mode == "Internship" else apprenticeships_collection).find(roles_query, {"role": 1})

        roles = sorted({doc["role"] for doc in roles_cursor})
        selected_role = st.selectbox("Select Job Role:", roles if roles else ["No roles available"])



    # Skills selection
    skill_options = [
        "Python", "Java", "C++", "C", "JavaScript", "TypeScript", "R", "Go", "Rust", "Kotlin", "Swift", "SQL",
        "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Flask", "Django", "Bootstrap", "Tailwind CSS",
        "Pandas", "NumPy", "Matplotlib", "Seaborn", "scikit-learn", "TensorFlow", "PyTorch", "Keras", "OpenCV",
        "NLP", "Deep Learning", "Machine Learning", "Computer Vision", "Data Visualization", "Data Cleaning",
        "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "CI/CD", "Terraform", "Jenkins", "Linux",
        "Shell Scripting", "Git", "GitHub", "Bitbucket", "MySQL", "PostgreSQL", "MongoDB", "Firebase",
        "Redis", "Oracle", "Hadoop", "Spark", "Hive", "Power BI", "Tableau", "Figma", "Canva", "Notion",
        "Jira", "Slack", "Heroku", "Streamlit", "Arduino", "Raspberry Pi", "Embedded C", "MATLAB", "Verilog",
        "VHDL", "IoT", "Sensors", "Robotics", "Nmap", "Wireshark", "Burp Suite", "Metasploit", "Ethical Hacking",
        "Network Security", "Communication", "Teamwork", "Leadership", "Problem Solving", "Critical Thinking",
        "Time Management", "UI/UX", "Graphic Design", "Technical Writing", "Agile", "Scrum", "Project Management",
        "Business Analysis","Excel"
    ]
    selected_skills = st.multiselect("Select Skills:", options=skill_options)

    # Location and mode handling
    selected_location = None
    internship_payment = None
    stipend_type = None
    if mode == "Internship":
        internship_mode = st.selectbox("Select Internship Type:", ["Online", "Offline"])
        if internship_mode == "Offline":
            selected_location = st.selectbox("Select Location:", indian_cities)
        internship_payment = st.selectbox("Select Stipend Type:", ["Paid", "Unpaid"])
    elif mode == "Apprenticeship":
        selected_location = st.selectbox("Select Location:", indian_cities)
        stipend_type = st.selectbox("Select Stipend Type:", ["Fixed", "Performance-Based", "Unpaid", "Paid"])
    if "search_performed" not in st.session_state:
        st.session_state.search_performed = False
   
   

    # Search button
    # Determine which collection to use upfront
    mode = st.session_state.get("internship_mode", "Internship")
    collection = internships_collection if mode == "Internship" else apprenticeships_collection

    if st.button(f"Find {mode}s"):
        st.session_state.search_performed = True

        # Build the base query dictionary dynamically
        query = {}

        # Filter by type: exact match case-insensitive (Internship or Apprenticeship)
        query["type"] = {"$regex": f"^{mode}$", "$options": "i"}

        # Sector (Government or Private)
        if internship_type:
            query["sector"] = {"$regex": f"^{internship_type}$", "$options": "i"}

        # Category (Hardware or Software)
        if category:
            query["category"] = {"$regex": f"^{category}$", "$options": "i"}

        # Mode filter (Online/Offline for Internship only)
        if mode == "Internship" and internship_mode:
            query["mode"] = {"$regex": f"^{internship_mode}$", "$options": "i"}

        # Stipend filter depending on mode
        stipend_value = internship_payment if mode == "Internship" else stipend_type
        if stipend_value:
            query["stipend"] = {"$regex": f"^{stipend_value}$", "$options": "i"}

        # Role filter, ignore "No roles available"
        if selected_role and selected_role != "No roles available":
            query["role"] = {"$regex": f"^{selected_role}$", "$options": "i"}

        # Location filter (optional)
        if selected_location:
            query["location"] = {"$regex": f"^{selected_location}$", "$options": "i"}

        # Skills filter: MongoDB expects at least one matching skill in the "skills" array
        if selected_skills and len(selected_skills) > 0:
            query["skills"] = {"$in": selected_skills}

        # Query MongoDB
        results = list(collection.find(query))
        
        if results:
            df = pd.DataFrame(results)

            # Calculate skill match scores only if skills selected and "skills" field present
            if selected_skills and "skills" in df.columns:
                skills_string = ", ".join(selected_skills)
                skill_matrix = TfidfVectorizer().fit_transform(
                    [", ".join(s) for s in df["skills"]] + [skills_string]
                )
                scores = cosine_similarity(skill_matrix[-1], skill_matrix[:-1]).flatten()
                df["Match Score"] = scores
                df = df.sort_values(by="Match Score", ascending=False)

            st.session_state.filtered_df = df

            if len(selected_skills) < 2:
                st.warning("⚠️ Please select at least two skills to find relevant listings.")
                st.session_state.filtered_df = None

        else:
            st.warning("❌ No matching listings found. Showing one sample listing from DB:")
            sample = collection.find_one()
            st.write(sample)
            st.session_state.filtered_df = None

            if len(selected_skills) < 2:
                st.warning("⚠️ Please select at least two skills to find internships.")
                st.session_state.filtered_df = None
            else:
                collection = internships_collection if mode == "Internship" else apprenticeships_collection
                roles_cursor = collection.find(roles_query, {"role": 1})
                roles = sorted({doc["role"] for doc in roles_cursor if "role" in doc})
                stipend_filter = internship_payment if mode == "Internship" else stipend_type
                if selected_role == "No roles available":
                    selected_role = None  
                
                
                    st.session_state.filtered_df = df
                    st.write("🔒 Logged in as:", st.session_state.get("user_email", "No email found"))
                else:
                    st.session_state.filtered_df = None
                    st.warning("No matching listings found.")
                    
        
            
    df=st.session_state.get("filtered_df") # Show results
    if st.session_state.get("search_performed") and df is not None and not df.empty:
        st.subheader(f"Recommended {mode}s:")

        if "applied_roles" not in st.session_state:
            st.session_state.applied_roles = []

        for idx, row in df.iterrows():
            with st.expander(f"📌 {row['role']} at {row['company']}"):
                st.markdown(f"""
                **Company:** {row.get('company', 'N/A')}  
                **Role:** {row.get('role', 'N/A')}  
                **Skills Required:** {', '.join(row.get('skills', []))}  
                **Stipend:** {row.get('stipend', 'N/A')}  
                **Location:** {row.get('location', 'N/A')}  
                **Apply By:** {row.get('apply_by', 'N/A')}  
                """)

                apply_key = f"apply_{row['_id']}"

                # Only show the button if the user hasn't applied yet
                if apply_key not in st.session_state.applied_roles:
                    clicked = st.button(f"Apply", key=apply_key)

                    # If the user clicks "Apply", add the role to the applied roles list
                    if clicked:
                        st.session_state.applied_roles.append(apply_key)
                        st.success(f"✅ Application submitted for {row['role']} at {row['company']}")

                        # Check if user email exists in session state
                        user_email = st.session_state.get("user_email", "")
                        if user_email:
                            # Send the email and show success/failure
                            result = send_application_email(user_email, row['role'], row['company'])
                            if result:
                                st.info("📧 Email sent successfully!")
                            else:
                                st.error("❌ Email failed to send")
                        else:
                            st.error("❌ No user email in session state")
                else:
                    st.info("ℹ️ Already applied for this role.")
    else:
        if st.session_state.get("search_performed"):
            st.warning("No matching listings found.")
        else:
            message = f"Use the filters above and click Find {mode}s"
            st.info(message)
if not st.session_state.logged_in:
    if st.session_state.show_signup:
        show_signup()
    else:
        show_login()
else:
    if st.session_state.show_profile:
        show_profile()
    elif st.session_state.page == "internexus":
            show_internexus()
    elif st.session_state.page == "main":
            show_main_page()
    if st.sidebar.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.page = "login"
        st.session_state.show_profile = False
        st.session_state.show_signup = False
        st.session_state.user_email = ""
        st.session_state.filtered_df = None
        st.session_state.applied_roles = []
        st.rerun()